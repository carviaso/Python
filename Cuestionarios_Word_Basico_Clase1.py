from docx import Document

# Crear documento Word para cuestionarios del curso completo
doc = Document()
doc.add_heading('Curso de Microsoft Word Básico', 0)
doc.add_heading('Cuestionarios de Validación de Conocimientos', level=1)

# Datos por clase
clases = {
    "Clase 1: Interfaz de Word": [
        ("¿Cuál es la función de la cinta de opciones en Word?", ["a) Guardar el documento", "b) Mostrar herramientas organizadas en pestañas", "c) Imprimir el documento", "d) Escribir comentarios"]),
        ("¿Dónde se encuentra el botón para guardar un documento?", ["a) Pestaña Inicio", "b) Barra de herramientas de acceso rápido", "c) Área de trabajo", "d) Pestaña Insertar"]),
        ("La barra donde puedes cambiar el tipo y tamaño de fuente se llama:", ["a) Barra de estado", "b) Panel de navegación", "c) Grupo de fuente", "d) Barra de herramientas"]),
        ("¿Qué función tiene la vista previa del documento?", ["a) Insertar gráficos", "b) Ver cómo se verá el documento al imprimirlo", "c) Crear una copia", "d) Cambiar idioma"]),
        "Verdadero o Falso: La cinta de opciones se puede minimizar.",
        "Verdadero o Falso: La barra de estado muestra número de páginas y palabras.",
        "Verdadero o Falso: La pestaña Diseño permite insertar imágenes.",
        "Verdadero o Falso: Puedes personalizar la barra de herramientas de acceso rápido.",
        "Menciona tres elementos que puedes encontrar en la interfaz de Word.",
        "¿Para qué sirve la pestaña Archivo?"
    ]
}

# Agregar contenido por clase
for clase, preguntas in clases.items():
    doc.add_page_break()
    doc.add_heading(clase, level=1)
    for item in preguntas:
        if isinstance(item, tuple):
            pregunta, opciones = item
            doc.add_paragraph(pregunta, style='List Number')
            for op in opciones:
                doc.add_paragraph(op, style='List Bullet')
        else:
            doc.add_paragraph(item, style='List Number')

# Guardar el archivo
file_path = "d:/Cuestionarios_Word_Basico_Clase1.docx"
doc.save(file_path)
file_path
