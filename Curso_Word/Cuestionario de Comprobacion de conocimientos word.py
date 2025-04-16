# -*- coding: utf-8 -*-
"""
Created on Tue Apr 15 12:32:28 2025

@author: carvi
"""

from docx import Document

# Crear documento Word para cuestionarios del curso completo
doc = Document()
doc.add_heading('Curso de Microsoft Word Básico', 0)
doc.add_heading('Cuestionarios de Validación de Conocimientos', level=1)

# Datos por clase
clases = {
    "Clase 1: Interfaz de Word": [
        ("¿Cuál es la función de la cinta de opciones en Word?", ["a) Guardar el documento", "b) Mostrar herramientas organizadas en pestañas", "c) Imprimir el documento", "d) Escribir comentarios"]),
        ("¿Dónde se encuentra el botón para guardar un documento?", ["a) Pestaña Inicio", "b) Barra de herramientas de acceso rápido", "c) Área de trabajo", "d) Pestaña Insertar"]),
        ("La barra donde puedes cambiar el tipo y tamaño de fuente se llama:", ["a) Barra de estado", "b) Panel de navegación", "c) Grupo de fuente", "d) Barra de herramientas"]),
        ("¿Qué función tiene la vista previa del documento?", ["a) Insertar gráficos", "b) Ver cómo se verá el documento al imprimirlo", "c) Crear una copia", "d) Cambiar idioma"]),
        "Verdadero o Falso: La cinta de opciones se puede minimizar.",
        "Verdadero o Falso: La barra de estado muestra número de páginas y palabras.",
        "Verdadero o Falso: La pestaña Diseño permite insertar imágenes.",
        "Verdadero o Falso: Puedes personalizar la barra de herramientas de acceso rápido.",
        "Menciona tres elementos que puedes encontrar en la interfaz de Word.",
        "¿Para qué sirve la pestaña Archivo?"
    ],
    "Clase 2: Crear y editar documentos": [
        ("¿Qué combinación de teclas permite guardar un documento rápidamente?", ["a) Ctrl + S", "b) Ctrl + V", "c) Ctrl + X", "d) Ctrl + P"]),
        ("¿Qué opción utilizas para abrir un archivo existente?", ["a) Archivo > Imprimir", "b) Archivo > Abrir", "c) Inicio > Abrir", "d) Vista > Archivo"]),
        "Verdadero o Falso: Puedes eliminar texto con la tecla Suprimir.",
        "Verdadero o Falso: Word no permite trabajar con varios documentos a la vez.",
        "Describe los pasos para crear un nuevo documento.",
        "¿Cómo puedes cerrar un documento sin cerrar Word?"
    ],
    "Clase 3: Formato de texto y párrafos": [
        ("¿Qué botón usas para poner el texto en negrita?", ["a) N", "b) B", "c) I", "d) S"]),
        ("¿Cuál es el propósito del interlineado?", ["a) Cambiar el tamaño del texto", "b) Agregar espacios entre párrafos", "c) Separar líneas del texto", "d) Justificar el texto"]),
        "Verdadero o Falso: Se puede aplicar formato a varias palabras a la vez.",
        "Verdadero o Falso: La alineación derecha es la predeterminada en Word.",
        "Explica cómo cambiar el color de una palabra.",
        "Menciona dos formas de aplicar sangría a un párrafo."
    ],
    "Clase 4: Páginas y secciones": [
        ("¿Cómo insertas un salto de página?", ["a) Diseño > Saltos > Página", "b) Insertar > Página en blanco", "c) Diseño > Margen", "d) Archivo > Nueva página"]),
        ("¿Dónde puedes cambiar el tamaño de la página?", ["a) Insertar > Tamaño", "b) Diseño > Tamaño", "c) Archivo > Propiedades", "d) Inicio > Fuente"]),
        "Verdadero o Falso: Puedes cambiar orientación de una sola página con secciones.",
        "Verdadero o Falso: Márgenes se configuran desde la pestaña Diseño.",
        "Explica cuándo usar un salto de sección.",
        "Menciona los pasos para cambiar los márgenes de un documento."
    ],
    "Clase 5: Elementos gráficos": [
        ("¿Cómo insertas una imagen desde tu computadora?", ["a) Archivo > Abrir", "b) Insertar > Imagen", "c) Inicio > Imagen", "d) Vista > Insertar"]),
        ("¿Qué herramienta ajusta el texto alrededor de una imagen?", ["a) Ajuste de párrafo", "b) Insertar tabla", "c) Ajuste de texto", "d) Revisión"]),
        "Verdadero o Falso: Se puede cambiar el tamaño de una imagen manualmente.",
        "Verdadero o Falso: Las formas no se pueden personalizar.",
        "Describe cómo insertar una tabla.",
        "Explica cómo agregar y editar una forma."
    ],
    "Clase 6: Imágenes y gráficos": [
        ("¿Qué tipo de gráfico se usa para mostrar cambios en el tiempo?", ["a) Circular", "b) Barra", "c) Línea", "d) Columna"]),
        ("¿Cómo se cambia el diseño de un gráfico?", ["a) Inicio > Estilos", "b) Herramientas de gráfico > Diseño", "c) Insertar > Cambiar", "d) Archivo > Diseño"]),
        "Verdadero o Falso: Word permite insertar gráficos sin Excel.",
        "Verdadero o Falso: Las imágenes se pueden recortar en Word.",
        "Explica cómo insertar una imagen y ajustarla al texto.",
        "Menciona los pasos para insertar un gráfico de barras."
    ],
    "Clase 7: Revisión y corrección": [
        ("¿Qué tecla rápida activa el corrector ortográfico?", ["a) F7", "b) Ctrl + R", "c) Alt + F", "d) Ctrl + C"]),
        ("¿Qué permite el control de cambios?", ["a) Aumentar el zoom", "b) Ver modificaciones", "c) Imprimir", "d) Insertar tablas"]),
        "Verdadero o Falso: Puedes agregar comentarios en Word.",
        "Verdadero o Falso: Word no corrige gramática.",
        "Describe el proceso para usar el corrector ortográfico.",
        "¿Qué utilidad tiene agregar comentarios?"
    ],
    "Clase 8: Impresión y exportación": [
        ("¿Qué opción permite guardar en PDF?", ["a) Guardar como > PDF", "b) Exportar > JPEG", "c) Imprimir > Imagen", "d) Inicio > PDF"]),
        ("¿Qué permite la vista preliminar?", ["a) Cambiar fuente", "b) Ver apariencia antes de imprimir", "c) Insertar una tabla", "d) Crear índice"]),
        "Verdadero o Falso: Puedes imprimir páginas específicas.",
        "Verdadero o Falso: Word no permite exportar a otros formatos.",
        "Menciona las opciones básicas del menú Imprimir.",
        "Explica cómo guardar un documento en formato PDF."
    ],
    "Clase 9: Tablas y contenido estructurado": [
        ("¿Qué pestaña permite insertar una tabla?", ["a) Archivo", "b) Insertar", "c) Vista", "d) Revisión"]),
        ("¿Cómo se agrega una fila en una tabla?", ["a) Menú Insertar > Fila", "b) Botón derecho > Insertar fila", "c) Inicio > Fila", "d) Diseño > Nueva fila"]),
        "Verdadero o Falso: Las tablas no permiten formato.",
        "Verdadero o Falso: Puedes combinar celdas en una tabla.",
        "Describe cómo insertar una tabla con 3 filas y 2 columnas.",
        "Explica cómo dar formato a una tabla insertada."
    ],
    "Clase 10: Tabla de contenido": [
        ("¿Qué función tiene una tabla de contenido?", ["a) Agregar imágenes", "b) Mostrar títulos y su ubicación", "c) Crear una portada", "d) Insertar gráficos"]),
        ("¿Qué requisito se necesita para que Word genere una tabla de contenido?", ["a) Imágenes con texto alternativo", "b) Títulos con estilos", "c) Texto en columnas", "d) Documento guardado como PDF"]),
        "Verdadero o Falso: Las tablas de contenido se actualizan automáticamente.",
        "Verdadero o Falso: Word no permite personalizar el formato de la tabla de contenido.",
        "Explica cómo insertar una tabla de contenido automática.",
        "¿Cómo actualizas la tabla de contenido luego de modificar el documento?"
    ]
}

# Agregar contenido por clase
for clase, preguntas in clases.items():
    doc.add_page_break()
    doc.add_heading(clase, level=1)
    for item in preguntas:
        if isinstance(item, tuple):
            pregunta, opciones = item
            doc.add_paragraph(pregunta, style='List Number')
            for op in opciones:
                doc.add_paragraph(op, style='List Bullet')
        else:
            doc.add_paragraph(item, style='List Number')

# Guardar el archivo
file_path = "d:/Cuestionarios_Word_Basico.docx"
doc.save(file_path)
file_path
